import numpy as np
from docx import Document
from pandas import DataFrame

def docx_writer(num_people: int, dataframe: DataFrame, output_path = None):
    """
    num_people: number of people the dataframe is to be divided among
    dataframe: dataframe containing narratives to be evaluated
    path(optional): path to store the docx files 
    """

    chunks = np.array_split(dataframe, num_people)
    
    for index, chunk in enumerate(chunks):
        doc = Document()
        for _ in range(len(chunk)-1):
            doc.add_paragraph("Post id: " + str(chunk.iloc[_+1]["ids"]))
            doc.add_paragraph("Title: " + str(chunk.iloc[_+1]["title"])+ "\n")
            doc.add_paragraph("Distance: " + str(chunk.iloc[_+1]['distance'])+ "\n")

            try: 
                para = doc.add_paragraph(chunk.iloc[_+1]['selftext'])
                para.keep_together = True
                doc.add_page_break() 
            except:
                pass
            
        if  output_path != None:   
            doc.save(f"{output_path}\file_{index}.docx")

        else: 
            doc.save(f"file_{index}.docx")

def main(): 
    import pandas as pd 

    results_df = pd.read_csv(r"C:\Users\nlplab\Research\Data-Collection\notebooks\results_tokenized_full.csv")
    results_df = results_df[1:]
    results_df.reset_index(drop=True)

    docx_writer(num_people=5, dataframe = results_df)

if __name__ == "__main__": 
    main()